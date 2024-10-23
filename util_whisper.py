import whisper
from docx import Document
doc = Document()


model = whisper.load_model("base")
result = model.transcribe("2024-08-23 14-31-03.mp3", verbose=True, word_timestamps=True, prepend_punctuations="\"'“¿([{-" ,append_punctuations="\"'.。,，!！?？:：”)]}、")
print(result["segments"])
doc.add_paragraph(result["text"])
# doc.add_paragraph(result["segments"])
doc.save('transcript2.docx')
